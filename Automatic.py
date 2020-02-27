# guide: https://www.youtube.com/watch?v=26vNgM_wSAE

import docx
from docx.shared import Pt, RGBColor
# from docx.enum.style import WD_STYLE_TYPE
# import os
import pandas as pd
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH


df = pd.read_excel(r'Listagem.xlsx') 

# AS PRÓXIMAS 3 LINHAS É PORQUE TEM UMA FUCKING BARRA ( / ) 
# NO CABEÇALHO DA COLUNA DA TABELA NO EXCEL
# E TAMBÉM PARA REMOVER ESPAÇOS
cols = df.columns
cols = cols.map(lambda x: x.replace('/', '') if isinstance(x, (str)) else x)
cols = cols.map(lambda x: x.replace(' ', '') if isinstance(x, (str)) else x)
cols = cols.map(lambda x: x.replace('.', '') if isinstance(x, (str)) else x)
df.columns = cols
# É ISTO
for index, row in df.iterrows():
    if row.Resultado == 'Normal':    
        doc = docx.Document('model1.docx')
        estilo1 = doc.styles['Heading 1']
        font_estilo1 = estilo1.font
        font_estilo1.name = 'Calibri'
        font_estilo1.size = Pt(14)
        font_estilo1.color.rgb = RGBColor(0, 0, 0)
        para1FontName = "Calibri"
        para1 = doc.add_paragraph()
        para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para1.style = doc.styles['Heading 1']
        para1Nome = para1.add_run('Paciente\t: '+row.Nome)
        para1Nome.font.name = para1FontName
        para1Nome.bold = True
        para1Nome.add_break()
        para1Idade = para1.add_run('Idade      \t: '+str(row.Idade))
        para1Idade.font.name = para1FontName
        para1Idade.bold = False
        para1Idade.add_break()
        para1Idade.add_break()
        datalaudo = row.Data
        if isinstance(datalaudo, datetime):
            datalaudo = datalaudo.strftime('%d/%m/%Y')
        para1Data = para1.add_run('Data        \t: '+datalaudo)
        para1Data.font.name = para1FontName
        para1Data.bold = False
        para1Data.add_break()
        para1Data.add_break()
        para1Categoria = para1.add_run('Categoria     : '+str(row.EmpresaConvênio))
        para1Categoria.font.name = para1FontName
        para1Categoria.bold = False
        para1Categoria.add_break()
        para1Motivo = para1.add_run('Motivo  \t: '+str(row.TipodeExame))
        para1Motivo.font.name = para1FontName
        para1Motivo.bold = False
        para1Motivo.add_break()
        para1Medicacao = para1.add_run('Medicação\t: '+str(row.Medicação))
        para1Medicacao.font.name = para1FontName
        para1Medicacao.bold = False
        para1Medicacao.add_break()
        para1Medicacao.add_break()
        para1Medicacao.add_break()
        para1Medico = para1.add_run('Médico(a) solicitante: Dr(a). '+str(row.Médico))
        para1Medico.font.name = para1FontName
        para1Medico.bold = True
        para1Medico.add_break()
        para2 = doc.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para2.style = doc.styles['Normal']
        para2Resultado = para2.add_run('Resultado:')
        para2Resultado.font.name = para1FontName
        para2Resultado.font.size = Pt(14)
        para2Resultado.bold = True
        para2texto = para2.add_run(' ELETRENCEFALOGRAMA DIGITAL e MAPEAMENTO CEREBRAL' +
                                'de vigília considerados normais para a faixa' +
                                'etária e as condições de realização do traçado.')
        para2texto.font.name = para1FontName
        para2texto.font.size = Pt(14)
        para2texto.bold = False
        para3texto = para2.add_run('                                            .')
        para3texto.font.name = para1FontName
        para3texto.font.size = Pt(14)
        para3texto.font.color.rgb = RGBColor(255, 255, 255)
        para3texto.bold = False
        for hue in range(6):
            para3texto.add_break()
        pEEG = doc.add_paragraph()
        pEEG.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rEEG = pEEG.add_run('Eletroencefalograma Digital')
        rEEG.font.name = 'Calibri'
        rEEG.font.size = Pt(14)
        rEEG.bold = True
        rEEG.underline = True
        rEEG.add_break()
        rEEG.add_break()
        pEEG.paragraph_format.page_break_before = True
        pCdTec = doc.add_paragraph()
        rCdTec = pCdTec.add_run('Condições técnicas		: boas.')
        rCdTec.font.size = Pt(14)
        rCdTec.font.name = 'Calibri'
        pColPa = doc.add_paragraph()
        rColPa = pColPa.add_run('Colaboração do(a) paciente 	: boa.')
        rColPa.font.size = Pt(14)
        rColPa.font.name = 'Calibri'
        pFMinMax = doc.add_paragraph()
        rFMinMax = pFMinMax.add_run('Ritmo dominante posterior alfa na frequência ' +
                                    'de {0} Hz, de {1} a {2} mV.'.format(row.Freq, row.VMin, row.VMax))
        rFMinMax.font.size = Pt(14)                         
        rFMinMax.font.name = 'Calibri'
        pFiller1 = doc.add_paragraph()
        rFiller1 = pFiller1.add_run('Bem regulado em amplitude e frequência.')
        rFiller1.font.size = Pt(14)
        rFiller1.font.name = 'Calibri'
        pFiller2 = doc.add_paragraph()
        rFiller2 = pFiller2.add_run('Simétrico e síncrono.')
        rFiller2.font.size = Pt(14)
        rFiller2.font.name = 'Calibri'
        pFiller3 = doc.add_paragraph()
        rFiller3 = pFiller3.add_run('Sinusoidal e complexo.')
        rFiller3.font.size = Pt(14)
        rFiller3.font.name = 'Calibri'
        pFiller4 = doc.add_paragraph()
        rfiller4 = pFiller4.add_run('Bloqueio à abertura dos olhos	: completo.')
        rfiller4.font.size = Pt(14)
        rfiller4.font.name = 'Calibri'
        pFiller5 = doc.add_paragraph()
        rFiller5 = pFiller5.add_run('Ritmos patológicos generalizados	:não se observam.')
        rFiller5.font.size = Pt(14)
        rFiller5.font.name = 'Calibri'
        pFiller6 = doc.add_paragraph()
        rFiller6 = pFiller6.add_run('Ritmos patológicos focais		:não se observam.')
        rFiller6.font.size = Pt(14)
        rFiller6.font.name = 'Calibri'
        pFiller7 = doc.add_paragraph()
        rFiller7 = pFiller7.add_run('O(A) paciente não entrou em sonolência.')
        rFiller7.font.size = Pt(14)
        rFiller7.font.name = 'Calibri'
        pFiller8 = doc.add_paragraph()
        rFiller8 = pFiller8.add_run('Hiperventilação (mais de 3 minutos) :' +
                                    'não modifica essencialmente o traçado.''')
        rFiller8.font.size = Pt(14)
        rFiller8.font.name = 'Calibri'      
        rFiller8.add_break()
        rFiller8.add_break()                          
        pMCeb = doc.add_paragraph()
        pMCeb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rMCeb = pMCeb.add_run('Mapeamento Cerebral')
        rMCeb.bold = True
        rMCeb.underline = True
        rMCeb.font.size = Pt(14)
        rMCeb.font.name = 'Calibri'
        rMCeb.add_break()
        rMCeb.add_break()
        pFiller9 = doc.add_paragraph()
        rfiller9 = pFiller9.add_run('Diagrama posicional \t: ')
        rfiller9.font.size = Pt(14)
        rfiller9.font.name = 'Calibri'
        pFiller10 = doc.add_paragraph()
        rFiller10 = pFiller10.add_run('Amplitude média \t: ')
        rFiller10.font.size = Pt(14)
        rFiller10.font.name = 'Calibri'
        pFiller11 = doc.add_paragraph()
        rFiller11 = pFiller11.add_run('Espectro dinâmico \t:ritmo ' +
                                    'dominante posterior na frequência de' +
                                    ' {} Hz.'.format(row.Freq))
        rFiller11.font.size = Pt(14)
        rFiller11.font.name = 'Calibri'
        pFiller12 = doc.add_paragraph()
        rFiller12 = pFiller12.add_run('Histograma\t\t\t:ritmo dominante posterior alfa.')
        rFiller12.font.size = Pt(14)
        rFiller12.font.name = 'Calibri'
        pFiller13 = doc.add_paragraph()
        rFiller13 = pFiller13.add_run('Relação alfa/teta\t\t: normal.')
        rFiller13.font.size = Pt(14)
        rFiller13.font.name = 'Calibri'
        pFiller14 = doc.add_paragraph()
        rFiller14 = pFiller14.add_run('Diferença espectral \t: ')
        rFiller14.font.size = Pt(14)
        rFiller14.font.name = 'Calibri'
        pFiller15 = doc.add_paragraph()
        rFiller15 = pFiller15.add_run('Mapa de frequências \t: ')
        rFiller15.font.size = Pt(14)
        rFiller15.font.name = 'Calibri'
        doc.save('Resultado/' + row.Nome + '.docx')
